"""
Document Parser - Extract text from various file formats
Supports: .txt, .docx, .pdf
"""

import os
from pathlib import Path
from typing import Optional

# PDF parsing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("  PDF libraries not installed. Install with: pip install PyPDF2 pdfplumber")

# Word document parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("  python-docx not installed. Install with: pip install python-docx")


class DocumentParser:
    """
    Extract text from various document formats.
    """
    
    def __init__(self):
        self.supported_formats = ['.txt', '.docx', '.pdf']
        
    def parse_file(self, file_path: str) -> Optional[str]:
        """
        Parse a file and extract text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text or None if parsing failed
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            print(f" File not found: {file_path}")
            return None
        
        # Get file extension
        extension = file_path.suffix.lower()
        
        # Route to appropriate parser
        if extension == '.txt':
            return self._parse_txt(file_path)
        elif extension == '.docx':
            return self._parse_docx(file_path)
        elif extension == '.pdf':
            return self._parse_pdf(file_path)
        else:
            print(f"  Unsupported file format: {extension}")
            print(f"   Supported formats: {self.supported_formats}")
            return None
            
    def _parse_txt(self, file_path: Path) -> Optional[str]:
        """Parse .txt file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text
        except Exception as e:
            print(f" Error reading .txt file: {e}")
            return None
    
    def _parse_docx(self, file_path: Path) -> Optional[str]:
        """Parse .docx file (Microsoft Word)."""
        if not DOCX_AVAILABLE:
            print(" python-docx not installed!")
            return None
        
        try:
            doc = Document(file_path)
            
            # Extract all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Skip empty paragraphs
                    paragraphs.append(text)
            
            # Extract text from tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            
            # Join all text
            full_text = '\n'.join(paragraphs)
            
            print(f" Extracted {len(full_text)} characters from .docx")
            return full_text
            
        except Exception as e:
            print(f" Error reading .docx file: {e}")
            return None
    
    def _parse_pdf(self, file_path: Path) -> Optional[str]:
        """Parse .pdf file."""
        if not PDF_AVAILABLE:
            print(" PDF libraries not installed!")
            return None
        
        # Try pdfplumber first (better text extraction)
        try:
            text = self._parse_pdf_with_pdfplumber(file_path)
            if text and len(text) > 100:
                return text
        except Exception as e:
            print(f"  pdfplumber failed: {e}")
        
        # Fallback to PyPDF2
        try:
            text = self._parse_pdf_with_pypdf2(file_path)
            if text and len(text) > 100:
                return text
        except Exception as e:
            print(f" PyPDF2 failed: {e}")
        
        return None
    
    def _parse_pdf_with_pdfplumber(self, file_path: Path) -> Optional[str]:
        """Parse PDF using pdfplumber (better for formatted PDFs)."""
        import pdfplumber
        
        pages_text = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
        
        full_text = '\n'.join(pages_text)
        print(f" Extracted {len(full_text)} characters from .pdf (pdfplumber)")
        return full_text
    
    def _parse_pdf_with_pypdf2(self, file_path: Path) -> Optional[str]:
        """Parse PDF using PyPDF2 (fallback method)."""
        import PyPDF2
        
        pages_text = []
        
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
        
        full_text = '\n'.join(pages_text)
        print(f" Extracted {len(full_text)} characters from .pdf (PyPDF2)")
        return full_text
    
    def parse_directory(self, directory: str, file_pattern: str = '*') -> dict:
        """
        Parse all supported files in a directory.
        
        Args:
            directory: Path to directory
            file_pattern: Glob pattern (e.g., '*.pdf', 'cv_*')
            
        Returns:
            Dict mapping filenames to extracted text
        """
        directory = Path(directory)
        results = {}
        
        if not directory.exists():
            print(f" Directory not found: {directory}")
            return results
        
        # Find all files matching pattern
        for ext in self.supported_formats:
            pattern = f"{file_pattern}{ext}"
            for file_path in directory.glob(pattern):
                print(f"\n Parsing: {file_path.name}")
                text = self.parse_file(file_path)
                if text:
                    results[file_path.name] = text
        
        return results


def test_parser():
    """Test the document parser with sample files."""
    parser = DocumentParser()
    
    print("=" * 70)
    print("DOCUMENT PARSER TEST")
    print("=" * 70)
    
    # Test with a sample directory
    test_dir = Path('data/raw/Medicine')
    
    if test_dir.exists():
        print(f"\n Testing with directory: {test_dir}")
        results = parser.parse_directory(test_dir)
        
        print(f"\n Successfully parsed {len(results)} files:")
        for filename, text in results.items():
            print(f"  - {filename}: {len(text)} characters")
    else:
        print(f"\n  Test directory not found: {test_dir}")
        print("   Create some test files and try again!")


if __name__ == "__main__":
    test_parser()