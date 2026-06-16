import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def compile_report():
    doc = Document()
    
    # Optional: Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    files = ['report_part1.txt', 'report_part2.txt', 'report_part3.txt']
    
    for idx, file in enumerate(files):
        with open(file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                    continue
                
                # Check for headers
                if line.startswith('FACULTY OF') or line.startswith('Bachelor’s') or line.startswith('Second Progress') or line.startswith('Smart Career') or line.startswith('Connecting Students'):
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.5
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(line)
                    run.bold = True
                    run.font.size = Pt(14)
                elif line[0].isdigit() and (line[1] == '.' or (len(line) > 2 and line[1].isdigit() and line[2] == '.')):
                    # Section Header like "1. Introduction" or "2.1 Needs"
                    h = doc.add_heading(line, level=1 if line.count('.') == 1 else 2)
                    h.paragraph_format.line_spacing = 1.5
                elif line.startswith('The Advocate Agent') or line.startswith('The Skeptic Agent') or line.startswith('The Judge Agent'):
                    h = doc.add_heading(line, level=3)
                    h.paragraph_format.line_spacing = 1.5
                elif line.startswith('-'):
                    # Bullet point
                    p = doc.add_paragraph(line, style='List Bullet')
                    p.paragraph_format.line_spacing = 1.5
                elif line.startswith('Period:') or line.startswith('Phase:') or line.startswith('Activities:') or line.startswith('Status:'):
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.5
                    run = p.add_run(line.split(':', 1)[0] + ':')
                    run.bold = True
                    if ':' in line:
                        p.add_run(line.split(':', 1)[1])
                elif line.startswith('FACULTY DEAN') or line.startswith('HEAD OF') or line.startswith('ACADEMIC ADVISOR') or line.startswith('BACHELOR’S CANDIDATE') or line.startswith('Tirana'):
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.5
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run(line)
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.line_spacing = 1.5
                    
        # Add page break between major sections (optional, but good for structure)
        # if idx < len(files) - 1:
        #     doc.add_page_break()

    output_path = 'Second_Progress_Report.docx'
    doc.save(output_path)
    print(f"Successfully created {output_path}")

if __name__ == "__main__":
    compile_report()
